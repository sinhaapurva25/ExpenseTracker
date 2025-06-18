from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_horizontal_line(paragraph):
    p = paragraph._element
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '000000')
    pBdr.append(bottom)
    pPr.append(pBdr)

def set_paragraph_format(paragraph, font_size=11, font_name='Calibri', indent=0, space_after=6):
    paragraph.style.font.name = font_name
    paragraph.style.font.size = Pt(font_size)
    paragraph.paragraph_format.left_indent = Inches(indent)
    paragraph.paragraph_format.space_after = Pt(space_after)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

doc = Document()

# Header: Name and Contact Info
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("APURVA SINHA")
run.bold = True
run.font.size = Pt(14)
run.font.name = 'Calibri'
p.paragraph_format.space_after = Pt(2)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("sinhaapurva25@gmail.com")
run.add_break()
run = p.add_run("GitHub")
run.add_break()
run = p.add_run("LinkedIn")
run.add_break()
run = p.add_run("+918420979565")
set_paragraph_format(p, font_size=11, space_after=10)

# Section: Experience
p = doc.add_heading("Experience", level=1)
set_paragraph_format(p, font_size=12, space_after=6)
add_horizontal_line(p)

# Standard Chartered Bank
p = doc.add_paragraph()
run = p.add_run("Standard Chartered Bank, GBS Pvt Ltd")
run.bold = True
run = p.add_run("                                                                                                              Bangalore, KA")
set_paragraph_format(p, indent=0)
p = doc.add_paragraph()
run = p.add_run("Specialist, Data Engineering")
run.bold = True
run = p.add_run("                                                                                                                            April 2023-Present")
set_paragraph_format(p, indent=0)

bullets = [
    "Developed real-time data pipelines using Scala, Spark Structured Streaming, and Kafka to ingest high-velocity data into HBase and Hive, achieving 228 TPS in a Spark-Iceberg performance benchmarking POC.",
    "Optimized Spark-based pipelines in Python and Scala, automating table creation in Hive and bypassing configuration steps, reducing processing time by 20%.",
    "Built testing frameworks using Python and Spark, creating simulator code for Kafka message testing and unit tests for Spark 3.2 and Hive 3 migrations during HDP to CDP transitions.",
    "Led migration efforts from HDP to CDP, replacing HDFS with ABFS in Spark and Scala codebases, and supported Hive and HBase deployments with regression testing.",
    "Created in-house tools using Python and Spark for code scanning (regex-based suggestions for Spark 3.2 upgrades) and HDFS Analyser Utility to compare HBase and Hive data during migrations.",
    "Automated workflows with Apache Airflow and Python, scheduling Spark and Hive jobs in the GDP project to enhance data pipeline efficiency.",
    "Enhanced data ingestion with Python and Spark Structured Streaming, implementing JSON tokenization and file housekeeping for Kafka-sourced data into Hive and HBase.",
    "Supported DevOps workflows by writing Python unit tests for Spark and Kafka components, resolving bugs in HBase and Hive pipelines across production and lower environments.",
    "Improved pipeline performance using Scala and Spark, optimizing Spark memory overhead for Hive and HBase processes in CDP/HDP environments.",
    "Increased pipeline flexibility using Scala, Spark, and Airflow, enabling configurable table suffixes and key tab paths in Hive and HBase, and bypassing ISD validators for Kafka-driven pipelines."
]
for bullet in bullets:
    p = doc.add_paragraph(bullet, style='ListBullet')
    set_paragraph_format(p, indent=0.5)

# Oracle Financial Services
p = doc.add_paragraph()
run = p.add_run("Oracle Financial Services Software Pvt Ltd")
run.bold = True
run = p.add_run("                                                                                                         Bangalore, KA")
set_paragraph_format(p, indent=0)
p = doc.add_paragraph()
run = p.add_run("Staff Consultant (Python Developer)")
run.bold = True
run = p.add_run("                                                                                                    March 2022-March 2023")
set_paragraph_format(p, indent=0)
p = doc.add_paragraph("Project: Rainier Beacon", style='Italic')
set_paragraph_format(p, indent=0.25)
p = doc.add_paragraph("Tech stacks: Python, Big Data, ETL, Shell Scripting, Unix, Hadoop, Hive, Scala, Apache Spark", style='Italic')
set_paragraph_format(p, indent=0.25)

bullets = [
    "Developed Python codes and shell scripts to auto load large sets of raw share market (trade category) data to data warehouse, thereby reducing manual entry rate by 90%.",
    "Automate the run and control of ETL flow and keep the pipeline up for the Business team.",
    "Optimized the Hive table design, which reduced storage use by 50%.",
    "Hadoop cluster migration - developed codes to save all table creation and files automatically, thereby reducing manual work by 95%.",
    "Regularly monitor the data loads to Hive and report any discrepancies and defects within the data to the BI Engineers.",
    "Mentored junior team members.",
    "Performed POC on reducing the time of existing ETL jobs and migrating the codebase from Java to Scala."
]
for bullet in bullets:
    p = doc.add_paragraph(bullet, style='ListBullet')
    set_paragraph_format(p, indent=0.5)

# 4C Tech Ventures
p = doc.add_paragraph()
run = p.add_run("4C Tech Ventures Pvt Ltd")
run.bold = True
run = p.add_run("                                                                                                                                                  Remote")
set_paragraph_format(p, indent=0)
p = doc.add_paragraph()
run = p.add_run("Software Developer")
run.bold = True
run = p.add_run("                                                                                                                                    June 2021-March 2022")
set_paragraph_format(p, indent=0)
p = doc.add_paragraph("Project: Optical Fruit Grading System", style='Italic')
set_paragraph_format(p, indent=0.25)
p = doc.add_paragraph("Tech stacks: Python, OpenCV, NumPy, Pandas", style='Italic')
set_paragraph_format(p, indent=0.25)

bullets = [
    "Used mathematics and predictive analysis to extract information from images at run time.",
    "Separate background from foreground to extract each separate object.",
    "Single handedly curated the image processing software for the project.",
    "Created high speed algorithms, providing a throughput of 5 fruits per sec."
]
for bullet in bullets:
    p = doc.add_paragraph(bullet, style='ListBullet')
    set_paragraph_format(p, indent=0.5)

# Zentron Labs
p = doc.add_paragraph()
run = p.add_run("Zentron Labs Pvt Ltd")
run.bold = True
run = p.add_run("                                                                                                                                             Bangalore, KA")
set_paragraph_format(p, indent=0)
p = doc.add_paragraph()
run = p.add_run("Associate Software Engineer")
run.bold = True
run = p.add_run("                                                                                                                       July 2019-April 2021")
set_paragraph_format(p, indent=0)
p = doc.add_paragraph()
run = p.add_run("Intern")
run.bold = True
run = p.add_run("                                                                                                                                                          August 2018-June 2019")
set_paragraph_format(p, indent=0)
p = doc.add_paragraph("Project: Optical Fruit Grading System", style='Italic')
set_paragraph_format(p, indent=0.25)
p = doc.add_paragraph("Tech stacks: Image Processing, LabVIEW, C++, JSON", style='Italic')
set_paragraph_format(p, indent=0.25)

bullets = [
    "Experimented with edge detection algorithms (Robert, Canny, Sobel, Prewitt, Shi-Tomasi) to get a camera's depth of field to auto-focus it.",
    "Devised a model to measure the depth of field of a camera in real-time, responding to live images, with 100% efficiency.",
    "Designed a standalone, in-house desktop app using LabVIEW to calibrate light and camera(s) of the imaging unit of a fruit sorting machine.",
    "Optimized and recreated a desktop app to train objects based on its color.",
    "Improved the user experience of the tool to create configuration files to help classify colored images of fruits.",
    "Implemented the code for classifying fruits based on color using the HSI color model.",
    "Documented the color classification model for inclusion in the knowledge base for team use.",
    "Used predictive analysis to segment out fruits/spheres using the RGB color planes, from a series of incoming images.",
    "Using statistics, developed an algorithm to track each fruit in the FOV as they travel from one end of the imaging unit to another.",
    "Using mathematics, developed algorithms to detect and separate multiple fruits in the field of view and make a decision based on the logic's output for further processing.",
    "Curated test cases (as images), debugged codes to get the needed output for tracking fruits.",
    "Added features in the project to root cause production issues.",
    "Improved the accuracy of the output of fruit size by 50% using a linear regression model.",
    "Collaborated with the product owner to stay current on the product features and the intended functionality of the algorithms."
]
for bullet in bullets:
    p = doc.add_paragraph(bullet, style='ListBullet')
    set_paragraph_format(p, indent=0.5)

# Section: Education
p = doc.add_heading("Education", level=1)
set_paragraph_format(p, font_size=12, space_after=6)
add_horizontal_line(p)

p = doc.add_paragraph()
run = p.add_run("Techno India University, Kolkata, WB")
run.bold = True
run = p.add_run("                                                                                                        July 2015-June 2019")
set_paragraph_format(p, indent=0)
p = doc.add_paragraph("COURSE: B. Tech in ECE. CGPA: 8.69.")
set_paragraph_format(p, indent=0.25)

# Section: Publication
p = doc.add_heading("Publication", level=1)
set_paragraph_format(p, font_size=12, space_after=6)
add_horizontal_line(p)

p = doc.add_paragraph()
run = p.add_run("IEEE")
run.bold = True
run = p.add_run("                                                                                                                                                                       November 2018")
set_paragraph_format(p, indent=0)
p = doc.add_paragraph("Fault Tolerant Architecture Design of a 4-bit Magnitude Comparator")
set_paragraph_format(p, indent=0.25)

# Save the document
doc.save("CV_APURVA_SINHA_Formatted.docx")